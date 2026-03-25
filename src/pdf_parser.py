"""文件文字擷取（支援 PDF / DOCX）"""

from pathlib import Path

import pdfplumber


def extract_text(filepath: str, max_pages: int | None = None) -> tuple[str, int]:
    """從 PDF 或 DOCX 擷取文字內容

    Returns:
        (text, page_count)
    """
    ext = Path(filepath).suffix.lower()

    if ext == ".docx":
        return _extract_docx(filepath)
    else:
        return _extract_pdf(filepath, max_pages)


def _extract_pdf(filepath: str, max_pages: int | None = None) -> tuple[str, int]:
    """從 PDF 擷取文字"""
    pages_text = []
    page_count = 0

    with pdfplumber.open(filepath) as pdf:
        page_count = len(pdf.pages)
        limit = max_pages if max_pages else page_count
        for i, page in enumerate(pdf.pages[:limit]):
            text = page.extract_text()
            if text:
                pages_text.append(text)

    full_text = "\n\n".join(pages_text)
    return full_text, page_count


def _extract_docx(filepath: str) -> tuple[str, int]:
    """從 DOCX 擷取文字"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安裝 python-docx: pip install python-docx")

    doc = Document(filepath)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也擷取表格內容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    full_text = "\n\n".join(paragraphs)
    # DOCX 沒有明確的頁數概念，用段落數估算
    estimated_pages = max(1, len(paragraphs) // 20)
    return full_text, estimated_pages
